"""
國文第四課〈歲月跟著〉段考試卷產生器
修正版：強制替換佈景主題東亞字體，解決 MS Mincho 亂碼問題
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

ZH_FONT  = '微軟正黑體'
EN_FONT  = 'Times New Roman'
LINE_SPC = 276   # 240 = single, 276 = 1.15x

# ══════════════════════════════════════════════════════════════════════════════
# 字體修正核心
# ══════════════════════════════════════════════════════════════════════════════

def fix_theme_east_asia(doc, zh_font=ZH_FONT):
    """將文件主題的 East Asia 字體改成指定中文字體，避免 MS Mincho 出現。"""
    for rel in doc.part.rels.values():
        if 'theme' in rel.reltype.lower():
            blob = rel.target_part.blob
            tree = etree.fromstring(blob)
            ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            # majorFont / minorFont 裡的 <a:ea typeface="...">
            for ea in tree.iter(f'{{{ns}}}ea'):
                ea.set('typeface', zh_font)
            rel.target_part._blob = etree.tostring(
                tree, xml_declaration=True, encoding='UTF-8', standalone=True)
            break

def _set_rFonts(rPr, zh=ZH_FONT, en=EN_FONT):
    """在 rPr 上設定字體，並移除 theme 字體屬性。"""
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    # 移除 theme 覆蓋屬性
    for attr in (qn('w:asciiTheme'), qn('w:eastAsiaTheme'),
                 qn('w:hAnsiTheme'),  qn('w:cstheme')):
        if attr in rFonts.attrib:
            del rFonts.attrib[attr]
    rFonts.set(qn('w:ascii'),    en)
    rFonts.set(qn('w:hAnsi'),    en)
    rFonts.set(qn('w:eastAsia'), zh)
    rFonts.set(qn('w:cs'),       zh)

def fix_styles(doc, zh=ZH_FONT, en=EN_FONT):
    """修正 docDefaults 與 Normal style 的字體設定。"""
    # docDefaults
    styles_el = doc.styles.element
    docDef = styles_el.find(qn('w:docDefaults'))
    if docDef is not None:
        rPrDef = docDef.find(qn('w:rPrDefault'))
        if rPrDef is not None:
            rPr = rPrDef.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                rPrDef.append(rPr)
            _set_rFonts(rPr, zh, en)
    # Normal style
    for style in doc.styles:
        if style.name == 'Normal':
            sel = style.element
            rPr = sel.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                sel.append(rPr)
            _set_rFonts(rPr, zh, en)
            break

# ══════════════════════════════════════════════════════════════════════════════
# 段落 / Run 工具
# ══════════════════════════════════════════════════════════════════════════════

def set_spacing(p, before_pt=3, after_pt=2):
    pPr = p._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    sp.set(qn('w:before'),    str(int(before_pt * 20)))
    sp.set(qn('w:after'),     str(int(after_pt  * 20)))
    sp.set(qn('w:line'),      str(LINE_SPC))
    sp.set(qn('w:lineRule'),  'auto')

def styled_run(p, text, size=12, bold=False, italic=False,
               color=None, zh=ZH_FONT, en=EN_FONT):
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    _set_rFonts(rPr, zh, en)
    return run

def add_para(doc, text='', size=12, bold=False, italic=False,
             color=None, indent=0, before=3, after=2,
             align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    set_spacing(p, before, after)
    if text:
        styled_run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p

def add_section_header(doc, text, bg='DEEAF1'):
    p = add_para(doc, text, size=13, bold=True,
                 color=RGBColor(0x1F,0x49,0x7D), before=6, after=3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg)
    pPr.append(shd)
    return p

def add_ans(doc, answer, explain=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    set_spacing(p, 1, 1)
    styled_run(p, '[答案] ' + answer, size=11, bold=True,
               color=RGBColor(0xC0,0x00,0x00))
    if explain:
        styled_run(p, '  ' + explain, size=10, italic=True,
                   color=RGBColor(0x1E,0x7A,0x1E))

# ══════════════════════════════════════════════════════════════════════════════
# 題目資料
# ══════════════════════════════════════════════════════════════════════════════

SITUATION = (
    "【情境導入】小明整理家中舊物，翻出一本家族相簿，"
    "相片裡有祖父年輕時壯碩的身影、父母的結婚照、自己嬰兒時的笑臉。"
    "看著跨越數十年的照片，小明想起了向陽的詩——歲月跟著。"
    "請根據課文內容，回答下列各題。"
)

TRUE_FALSE = [
    ("1.", "歲月跟著的作者向陽，本名林淇瀁，為臺灣南投人。",
     "○", "課文作者介紹明載：本名林淇瀁，臺灣南投人。"),
    ("2.", "本詩以秒針象徵老年、分針象徵青壯、時針象徵童少時期的生命節奏。",
     "X", "秒針快速→童少；分針穩重→青壯；時針緩慢→老年，敘述顛倒。"),
    ("3.", "「馳過了兒童粲亮的眼睛」中，「粲亮」描繪的是童年充滿光彩、天真活潑的神情。",
     "○", "粲亮意為光彩鮮明，正面描寫童年的明亮眼神。"),
    ("4.", "本詩「馳過、翻閱、躡走」三個動詞由快而慢，以層遞手法呈現生命各階段的遞進。",
     "○", "三詞速度逐漸放慢，對應童少→青壯→老年，為層遞法。"),
    ("5.", "「圓柔的鐘面是生命的枷」意指時間無情折磨人，表達作者對生命的悲觀絕望。",
     "X", "下文接「輪迴地繞，土底的種子正開始抽芽」，傳達的是生命輪迴、生生不息的積極意涵。"),
]

SINGLE = [
    ("1.", "小明翻到祖父年輕時的照片，想起課文的作者介紹。下列關於向陽的敘述，何者正確？",
     [("A","本名林淇瀁，出生於臺灣北投"),
      ("B","致力於十行詩創作，不擅使用閩南語寫作"),
      ("C","致力於十行詩創作，擅長閩南語寫作，曾獲國家文藝獎"),
      ("D","國立臺灣大學中文系博士，現任教於國立政治大學")],
     "C","向陽：南投人，十行詩，擅閩南語，曾獲國家文藝獎。"),

    ("2.", "小明看著嬰兒時期的照片，想到詩中哪個意象最能對應童少時期的生命特質？",
     [("A","緩慢的時針是貓的腳步音"),
      ("B","滴答的秒針是蹄的聲音"),
      ("C","圓柔的鐘面是生命的枷"),
      ("D","雍容地耕")],
     "B","秒針快速象徵童少；蹄的聲音對應馬蹄奔跑，充滿活力。"),

    ("3.", "「歲月跟著犁耙沉穩地耕，雍容的分針是犁的鋒刃」，此節描述的是哪個人生階段？",
     [("A","嬰兒時期"),("B","童少時期"),("C","青壯時期"),("D","老年時期")],
     "C","分針對應青壯；犁耙耕作象徵青壯年穩重努力的特質。"),

    ("4.", "詩中「雍容地耕」最能凸顯青壯年的哪種特質？",
     [("A","急躁衝動、蓬勃朝氣"),("B","溫和莊重、沉穩奮鬥"),
      ("C","遲緩無力、逐漸老去"),("D","孤獨寂寞、悵然若失")],
     "B","雍容：溫和莊重、從容不迫，描寫青壯年穩健耕耘的生命態度。"),

    ("5.", "小明看著祖父老年的照片，眼神朦朧溫柔。詩中哪句最能呼應這個畫面？",
     [("A","馳過了兒童粲亮的眼睛"),("B","翻閱著六月的綠色大地"),
      ("C","躡走了老人眼角的水霧"),("D","土底的種子正開始抽芽")],
     "C","躡走對應老年；「眼角的水霧」描繪老人眼神朦朧的意象。"),

    ("6.", "「馳過→翻閱→躡走」三個動詞的排列，最主要呈現了何種效果？",
     [("A","速度由慢而快，生命越老越充滿活力"),
      ("B","速度由快而慢，以層遞手法表現生命各階段的遞進"),
      ("C","三個動詞速度相同，只是不同動作種類"),
      ("D","強調時間循環不息，表達作者的樂觀態度")],
     "B","層遞法：由快至慢，對應童少→青壯→老年的生命節奏。"),

    ("7.", "「粗糙的掌紋是犁的鋒刃」中，「掌紋」象徵什麼？",
     [("A","老人手掌乾燥，需要保養"),
      ("B","青壯年努力耕耘、辛勤勞動所留下的歲月痕跡"),
      ("C","童年玩樂時留下的傷疤"),
      ("D","命運注定、無法改變")],
     "B","掌紋代表青壯年辛勤勞動的痕跡，呼應犁耙耕作的意象。"),

    ("8.", "下列關於本詩形式的描述，何者正確？",
     [("A","自由詩，字數與節數不限"),
      ("B","每行十字，每節四行，四節成篇，形式整齊"),
      ("C","仿古典詩格律，平仄押韻"),
      ("D","散文詩形式，不分行")],
     "B","本詩每行十字，每節四行，四節成篇，形式整齊。"),

    ("9.", "「圓柔的鐘面是生命的枷，輪迴地繞，土底的種子正開始抽芽」最能傳達何種意涵？",
     [("A","時間是苦難，生命終將消逝"),
      ("B","生命雖受時間束縛，卻在輪迴中不斷更新延續"),
      ("C","種子象徵下一代，表達對年輕人的期許"),
      ("D","十二月是終點，一切歸於寂滅")],
     "B","「輪迴」與「抽芽」共同傳達生命循環不息、生生不已的正面意涵。"),

    ("10.","根據課文賞析，下列說明何者有誤？",
     [("A","動詞「馳過、翻閱、躡走」以層遞手法呈現生命各階段"),
      ("B","全詩以時鐘三針及鐘面為骨架，貫穿四個生命階段"),
      ("C","「兒童粲亮的眼睛」和「六月的綠色大地」都象徵青壯時期"),
      ("D","向陽認為古典詩的形式可去蕪存菁，使新詩作品更加精鍊")],
     "C","「兒童粲亮的眼睛」搭配秒針，象徵童少；只有「六月的綠色大地」象徵青壯，C有誤。"),
]

FILL = [
    ("1.","歲月跟著的作者向陽，本名（ ① ），臺灣（ ② ）人，詩作具有濃厚的實驗與反省精神。",
     ["① 林淇瀁","② 南投"]),
    ("2.","本詩以「（ ③ ）的秒針」、「（ ④ ）的分針」、「（ ⑤ ）的時針」象徵時間快慢節奏。",
     ["③ 滴答","④ 雍容","⑤ 緩慢"]),
    ("3.","詩中描寫童少用「（ ⑥ ）過」，青壯用「（ ⑦ ）閱」，老年用「（ ⑧ ）走」，由快而慢形成層遞。",
     ["⑥ 馳","⑦ 翻","⑧ 躡"]),
    ("4.","「緩慢的時針是（ ⑨ ）的腳步音」，以動物比喻老年時光緩緩流逝；「粲亮」的「粲」注音為（ ⑩ ）。",
     ["⑨ 貓","⑩ ㄘㄢˋ"]),
]

SHORT = [
    ("1.","請說明本詩「秒針、分針、時針」分別對應人生的哪三個階段，並各舉詩中一個具體意象加以說明。（5分）",
     "秒針→童少（兒童粲亮的眼睛）；分針→青壯（粗糙的掌紋是犁的鋒刃）；時針→老年（老人眼角的水霧）。"),
    ("2.","層遞法指三種以上事物依比例關係逐層排列的修辭手法。試以「馳過→翻閱→躡走」為例，說明此手法如何表現生命的變化。（5分）",
     "三詞由快（馳過奔跑）至慢（躡走輕步），對應童少→青壯→老年，生命節奏逐漸沉穩，形成由快而慢的層遞效果。"),
    ("3.","試解釋「粗糙的掌紋是犁的鋒刃，雍容地耕」這兩句詩的意涵。（5分）",
     "「粗糙的掌紋」象徵青壯年辛勤勞動的痕跡；以犁的鋒刃比喻掌紋的堅實；「雍容地耕」表現青壯年沉穩踏實、從容奮進的生命態度。"),
    ("4.","詩末「土底的種子正開始抽芽」傳達了什麼樣的生命觀？請結合「輪迴地繞」加以說明。（5分）",
     "「輪迴地繞」指生命如時鐘循環不已，雖然個體生命會凋謝，但「種子抽芽」代表生命在輪迴中延續更新，傳達生生不息、積極樂觀的生命觀。"),
]

PASSAGE = [
    "【延伸閱讀】非馬〈樹・四季〉（節選）",
    "",
    "    春",
    "你還是一樣年青",
    "深深藏在皺紋",
    "心底",
    "",
    "    夏",
    "非一隻羽毛豐滿的鳥",
    "高瞻遠矚",
    "的季節",
    "",
    "    秋",
    "最後一片落葉",
    "卻捉",
    "捉襟下來",
    "",
    "    冬",
    "老人的",
    "將手去",
    "呼嘯一北風",
    "苦笑著",
]

LITERACY = [
    ("1.","非馬以「四季」象徵人生階段，下列配對何者最符合詩意？",
     [("A","春—老年，夏—青壯"),
      ("B","春—童少，夏—青壯，秋—中年，冬—老年"),
      ("C","春—青壯，冬—童少"),
      ("D","春—老年，冬—童少")],
     "B","詩中春（年青）→夏（豐滿遠矚）→秋（落葉）→冬（老人），對應童少→青壯→中年→老年。"),

    ("2.","比較歲月跟著與樹・四季，兩首詩共同以何種事物為主軸描繪生命流逝？",
     [("A","時鐘與動物"),("B","四季與自然景物"),("C","花草與河流"),("D","海洋與星空")],
     "B","兩詩皆以四季（三月/六月/九月/十二月；春夏秋冬）搭配自然意象描繪生命各階段。"),

    ("3.","小明想用歲月跟著中的意象描述自己目前的人生階段。請引用詩中至少一個詞語，說明你認為最符合的詩節及理由。（開放式，言之成理即可）（5分）",
     None,
     "（開放式）如引用「兒童粲亮的眼睛」說明自己仍處童少充滿活力；或引用「雍容的分針」表示進入青壯。言之成理、引用正確即可。",
    ),

    ("4.","歲月跟著以「土底的種子正開始抽芽」作結，樹・四季則以「冬/老人的/苦笑著」收尾。試比較兩首詩對「生命終點」的態度有何不同？（5分）",
     None,
     "向陽以「抽芽」傳達生命輪迴、生生不息的積極樂觀；非馬以老人「苦笑」面對生命終點，帶有無奈與豁達並存的複雜情感，相對哀婉。",
    ),
]

# ══════════════════════════════════════════════════════════════════════════════
# 建立試卷
# ══════════════════════════════════════════════════════════════════════════════

def build(is_answer=False):
    doc = Document()

    # 修正字體
    fix_theme_east_asia(doc, ZH_FONT)
    fix_styles(doc, ZH_FONT, EN_FONT)

    # 頁面 A4（Cm() = EMU，1cm = 360000 EMU）
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    for attr in ('top_margin','bottom_margin','left_margin','right_margin'):
        setattr(sec, attr, Cm(1.9))

    # 標題
    add_para(doc, "國中國文（四）第四課 歲月跟著 段考試卷",
             size=18, bold=True, color=RGBColor(0x1F,0x49,0x7D),
             before=0, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "年級：八年級   學期：下學期   滿分：100 分",
             size=11, before=0, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 情境
    add_para(doc, SITUATION, size=10, italic=True,
             color=RGBColor(0x55,0x55,0x55), indent=0.5, before=0, after=6)

    # ── 一、是非題 ──
    add_section_header(doc, "一、是非題（每題 2 分，共 10 分）")
    add_para(doc, "請於括號內填入「○」（正確）或「X」（錯誤）。",
             size=10, indent=0.5, before=0, after=2)
    for num, q, ans, exp in TRUE_FALSE:
        add_para(doc, num + "  ( " + ("    " if not is_answer else ans + "  ") + " )  " + q,
                 size=12, indent=0.5)
        if is_answer:
            add_ans(doc, ans, exp)

    add_para(doc, size=12)

    # ── 二、選擇題 ──
    add_section_header(doc, "二、單一選擇題（每題 3 分，共 30 分）")
    for num, q, opts, ans, exp in SINGLE:
        add_para(doc, num + "  " + q, size=12, indent=0.5)
        for code, text in opts:
            add_para(doc, "(" + code + ")  " + text, size=12, indent=1.5,
                     before=0, after=1)
        if is_answer:
            add_ans(doc, "(" + ans + ")", exp)

    add_para(doc, size=12)

    # ── 三、填充題 ──
    add_section_header(doc, "三、填充題（每格 2 分，共 20 分）")
    for num, q, ans_list in FILL:
        add_para(doc, num + "  " + q, size=12, indent=0.5)
        if is_answer:
            add_ans(doc, "  ".join(ans_list))

    add_para(doc, size=12)

    # ── 四、簡答題 ──
    add_section_header(doc, "四、簡答 / 解釋題（每題 5 分，共 20 分）")
    for num, q, ans in SHORT:
        add_para(doc, num + "  " + q, size=12, indent=0.5)
        if is_answer:
            add_ans(doc, ans)

    add_para(doc, size=12)

    # ── 五、素養閱讀題 ──
    add_section_header(doc, "五、素養閱讀題（每題 5 分，共 20 分）")
    for line in PASSAGE:
        add_para(doc, line, size=11, indent=1.0, before=0, after=0)

    add_para(doc, size=12)

    for item in LITERACY:
        num, q = item[0], item[1]
        opts   = item[2]
        ans    = item[3]
        exp    = item[4] if len(item) > 4 else ""
        add_para(doc, num + "  " + q, size=12, indent=0.5)
        if opts:
            for code, text in opts:
                add_para(doc, "(" + code + ")  " + text, size=12,
                         indent=1.5, before=0, after=1)
            if is_answer:
                add_ans(doc, "(" + ans + ")", exp)
        else:
            if is_answer:
                add_ans(doc, ans)

    return doc


q_doc = build(is_answer=False)
a_doc = build(is_answer=True)

OUT_Q = r"C:\Users\北興\Desktop\claude-ai\exam_questions_v3.docx"
OUT_A = r"C:\Users\北興\Desktop\claude-ai\exam_answers_v3.docx"
q_doc.save(OUT_Q)
a_doc.save(OUT_A)
print("DONE:", OUT_Q)
print("DONE:", OUT_A)
