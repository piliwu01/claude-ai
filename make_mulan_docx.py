from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# === 頁面設定 A4 直式 ===
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

def set_font(run, name='微軟正黑體', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def set_para_shading(para, fill_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)

def set_cell_color(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=(44,62,80)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    size = 18 if level==1 else (14 if level==2 else 12)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_body(doc, text, size=10.5, color=(44,44,44), bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

# === 標題區 ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('木蘭詩 迷因講義')
set_font(run, size=22, bold=True, color=(190,30,45))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run('高中國文第五課｜樂府詩')
set_font(run2, size=11, color=(100,100,100))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(10)
run3 = p3.add_run('女生不只能織布——她可以打爆敵軍再回家煮飯')
set_font(run3, size=10, color=(150,100,0), bold=True)

doc.add_paragraph('─' * 50)

# === 速覽表格 ===
add_heading(doc, '本課重點速覽', level=2, color=(30,100,160))

table1 = doc.add_table(rows=1, cols=3)
table1.style = 'Table Grid'
hdr = table1.rows[0].cells
for i, txt in enumerate(['項目', '內容', '一句人話']):
    hdr[i].text = txt
    set_cell_color(hdr[i], '2E4057')
    for run in hdr[i].paragraphs[0].runs:
        set_font(run, bold=True, color=(255,255,255))

rows_data = [
    ('朝代/時代', '北魏，南北朝民歌（確切年代不詳）', '很久以前，作者已讀不回'),
    ('文體', '樂府詩（民間歌謠）', '那時代的 IG Reels，大家傳唱的神曲'),
    ('主角', '花木蘭（女扮男裝代父從軍）', '柔弱外表 + 鋼鐵意志，無敵combo'),
    ('全詩結構', '六節，三大部分', '準備→打仗→回家，英雄故事公式'),
    ('核心主題', '孝行忠義、巾幗英雄、不慕名利', '孝順滿分、打仗滿分、還不想當官'),
    ('寫作手法', '詳略得當、大量疊字、對偶排比', '重要的詳寫，不重要的略過'),
]

colors = ['F0F4F8', 'FFFFFF']
for idx, (a, b, c) in enumerate(rows_data):
    row = table1.add_row().cells
    row[0].text = a
    row[1].text = b
    row[2].text = c
    bg = colors[idx % 2]
    for cell in row:
        set_cell_color(cell, bg)

for row in table1.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(7)
    row.cells[2].width = Cm(7)

doc.add_paragraph()

# === 各節課文 + 梗 ===

def add_section(doc, title, poem, meme_type, meme_items, meme_label, pill_text):
    add_heading(doc, title, level=2, color=(160,30,60))

    # 課文
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(poem)
    set_font(run, name='標楷體', size=10.5, color=(60,40,20))
    set_para_shading(p, 'FFF8EC')

    # 梗圖標題
    p_mt = doc.add_paragraph()
    p_mt.paragraph_format.space_before = Pt(8)
    p_mt.paragraph_format.space_after = Pt(4)
    r = p_mt.add_run(meme_label)
    set_font(r, size=10, bold=True, color=(100,50,150))

    if meme_type == 'drake':
        for item_text in meme_items:
            mp = doc.add_paragraph()
            mp.paragraph_format.space_after = Pt(3)
            mp.paragraph_format.left_indent = Cm(0.5)
            mr = mp.add_run(item_text)
            if item_text.startswith('X') or 'X ' in item_text[:3]:
                color = (180,30,30)
                bg = 'FFE0E0'
            elif item_text.startswith('V') or 'V ' in item_text[:3]:
                color = (30,130,60)
                bg = 'D4EDDA'
            elif '不' in item_text[:5]:
                color = (180,30,30)
                bg = 'FFE0E0'
            else:
                color = (30,130,60)
                bg = 'D4EDDA'
            set_font(mr, size=10.5, bold=True, color=color)
            set_para_shading(mp, bg)

    elif meme_type == 'brain':
        brain_colors = ['EEEEEE', 'E0D5F0', 'C5B5E0', 'F5E6B0']
        brain_text_colors = [(80,80,80),(60,40,100),(80,30,120),(150,80,0)]
        prefixes = ['小腦', '中腦', '大腦', '超腦']
        for i, text in enumerate(meme_items):
            mp = doc.add_paragraph()
            mp.paragraph_format.space_after = Pt(2)
            mp.paragraph_format.left_indent = Cm(0.5)
            mr = mp.add_run(f'[{prefixes[i]}] {text}')
            set_font(mr, size=10.5, bold=(i==3), color=brain_text_colors[i])
            set_para_shading(mp, brain_colors[i])

    elif meme_type == 'gru':
        gru_colors = ['F0F4F8','E8EFF8','FFE8E8','FFF3CD']
        for i, text in enumerate(meme_items):
            mp = doc.add_paragraph()
            mp.paragraph_format.space_after = Pt(2)
            mp.paragraph_format.left_indent = Cm(0.5)
            mr = mp.add_run(f'Step {i+1}：{text}')
            set_font(mr, size=10.5, bold=(i==2), color=(44,62,80))
            set_para_shading(mp, gru_colors[i])

    elif meme_type == 'bf':
        bf_table = doc.add_table(rows=2, cols=3)
        bf_table.style = 'Table Grid'
        headers = ['男子（木蘭）', '紅衣女（女裝）', '被拋棄（男裝）']
        bf_colors = ['D5E8FF','FFD5D5','E0E0E0']
        hdr_row = bf_table.rows[0].cells
        for ci, h in enumerate(headers):
            hdr_row[ci].text = h
            set_cell_color(hdr_row[ci], bf_colors[ci])
            for run in hdr_row[ci].paragraphs[0].runs:
                set_font(run, size=9.5, bold=True, color=(44,62,80))
        content_row = bf_table.rows[1].cells
        for ci, text in enumerate(meme_items):
            content_row[ci].text = text
            set_cell_color(content_row[ci], bf_colors[ci])
            for run in content_row[ci].paragraphs[0].runs:
                set_font(run, size=10, color=(44,62,80))

    # 記憶口訣
    pill_p = doc.add_paragraph()
    pill_p.paragraph_format.space_before = Pt(8)
    pill_p.paragraph_format.space_after = Pt(10)
    pill_p.paragraph_format.left_indent = Cm(0.3)
    pr = pill_p.add_run('記憶口訣　' + pill_text)
    set_font(pr, size=10, color=(120,60,0))
    set_para_shading(pill_p, 'FFF3CD')

# 第一節
add_section(doc,
    '第一節：織布少女的危機',
    '唧唧復唧唧，木蘭當戶織。\n不聞機杼聲，惟聞女嘆息。\n昨夜見軍帖，可汗大點兵，\n軍書十二卷，卷卷有爺名。\n阿爺無大兒，木蘭無長兄，\n願為市鞍馬，從此替爺征。',
    'drake',
    [
        '不要 — 在家乖乖織布，讓老爸拖著病體上戰場',
        '要的 — 自己買馬買鞍，女扮男裝代父從軍',
    ],
    'Drake 梗：木蘭做決定的那一刻',
    '嘆→看→驚→決→GO！\n（嘆息→見軍帖→驚覺爺名→決定代父→出征）'
)

# 第二節
add_section(doc,
    '第二節：東南西北全跑透透',
    '東市買駿馬，西市買鞍韉，\n南市買轡頭，北市買長鞭。\n旦辭爺孃去，暮宿黃河邊，\n不聞爺孃喚女聲，但聞黃河流水鳴濺濺。',
    'brain',
    [
        '去一個市場買完所有東西',
        '東西南北四個市場分開買（帶動地方經濟）',
        '用四個市場寫出出征的鄭重和急迫',
        '其實是誇飾法！展現出征規格之高，非真的跑四市場！',
    ],
    'Expanding Brain 梗：疊字誇飾的意義',
    '東→駿馬 西→鞍韉 南→轡頭 北→長鞭\n記法：馬背著鞍，鞍配轡，轡頭拉馬，長鞭催！'
)

# 第三節
add_section(doc,
    '第三節：十年沙場（作者略過了）',
    '萬里赴戎機，關山度若飛。\n朔氣傳金柝，寒光照鐵衣。\n將軍百戰死，壯士十年歸。',
    'gru',
    [
        '詩要敘述木蘭十年征戰的英勇故事',
        '把十年的一切壓縮在六句話裡說完',
        '把十年的一切壓縮在六句話裡說完（再看一遍）',
        '老師說：這就是「詳略得當的敘事筆法」！',
    ],
    'Gru 計畫梗：詳略筆法的精髓',
    '「將軍百戰死，壯士十年歸」對偶句\n寫戰爭殘酷與時間之長——考試必考！'
)

# 第四節
add_section(doc,
    '第四節：不要當官！送我回家！',
    '歸來見天子，天子坐明堂。\n策勳十二轉，賞賜百千強。\n可汗問所欲，「木蘭不用尚書郎，\n願借明駝千里足，送兒還故鄉。」',
    'drake',
    [
        '不要 — 領尚書郎官職，從此在朝廷享盡榮華富貴',
        '要的 — 借一隻好駱駝，立刻回家陪爸媽吃飯',
    ],
    'Drake 梗：木蘭連 offer 都不收',
    '「不用尚書郎」→ 不慕名利\n「願借明駝千里足，送兒還故鄉」→ 孝行思念\n這兩句是表現木蘭性格的核心句！'
)

# 第五節
add_section(doc,
    '第五節：全家出動，弟弟磨刀殺豬歡迎',
    '爺孃聞女來，出郭相扶將。\n阿姊聞妹來，當戶理紅妝。\n小弟聞姊來，磨刀霍霍向豬羊。\n脫我戰時袍，著我舊時裳，\n當窗理雲鬢，對鏡貼花黃。',
    'bf',
    [
        '十年戰袍說脫就脫，換回女裝',
        '雲鬢、花黃、舊時裳——女兒本色',
        '戰時袍，功成身退，say goodbye',
    ],
    'Distracted BF 梗：木蘭人設大反轉',
    '爸媽→出城接（扶著走）\n阿姊→門口化妝（理紅妝）\n小弟→磨刀炒菜（霍霍聲）\n記法：「接、妝、刀」三字決'
)

# 第六節
add_section(doc,
    '第六節：大揭曉！兔子說明一切',
    '出門看火伴，火伴皆驚惶：\n「同行十二年，不知木蘭是女郎。」\n雄兔腳撲朔，雌兔眼迷離；\n兩兔傍地走，安能辨我是雄雌？',
    'brain',
    [
        '木蘭是個普通女生',
        '木蘭是個會打仗的女生',
        '木蘭是個讓全軍十二年認不出的女生',
        '「兩兔傍地走，安能辨我是雄雌」— 千古絕句收尾！',
    ],
    'Expanding Brain 梗：比喻的哲學',
    '雄兔（撲朔）+ 雌兔（迷離）兩兔同跑→分不清\n→比喻木蘭在戰場與男子無異\n→巾幗不讓鬚眉！（考試必考比喻意義）'
)

doc.add_paragraph('─' * 50)

# === 地雷區 ===
add_heading(doc, '考試地雷區', level=2, color=(180,30,30))

dangers = [
    ('地雷① 「唧唧復唧唧」是木蘭嘆氣聲', '「唧唧」是織布機的聲音，不是嘆氣！木蘭的嘆息是「惟聞女嘆息」才出現的。'),
    ('地雷② 東西南北四市是真的四個市集', '這是誇飾寫法，強調出征之鄭重，並非真的跑了四個市場。'),
    ('地雷③ 「策勳十二轉」十二是實數', '「十二」是虛數，代表最高等級的升遷獎賞，不是真的升了十二次官。'),
    ('地雷④ 木蘭最後選擇當尚書郎', '木蘭拒絕了尚書郎，選擇回家！這是「不慕名利」的關鍵情節。'),
    ('地雷⑤ 「兩兔傍地走」是說兔子跑法奇怪', '這是比喻：比喻木蘭在戰場上與男子無異，強調巾幗英雄形象。'),
]

for title, content in dangers:
    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(4)
    dp.paragraph_format.left_indent = Cm(0.3)
    dr1 = dp.add_run(title + '　')
    set_font(dr1, size=10.5, bold=True, color=(180,30,30))
    dr2 = dp.add_run(content)
    set_font(dr2, size=10.5, color=(60,30,30))
    set_para_shading(dp, 'FFF0F0')

doc.add_paragraph()

# === 速查表 ===
add_heading(doc, '考試救命速查表', level=2, color=(30,100,160))

lookup_table = doc.add_table(rows=1, cols=2)
lookup_table.style = 'Table Grid'
lookup_data = [
    ('作者', '佚名（北魏民間，後代文人潤色）'),
    ('文體特色', '雜言（以五言為主）、押韻自由可換韻、語言質樸'),
    ('樂府詩名稱', '漢武帝設「樂府」官署採集民歌，其歌辭稱樂府詩'),
    ('重要疊字', '唧唧、濺濺、啾啾、霍霍（狀聲詞，增強音樂性）'),
    ('詳略筆法', '詳：準備、離別、返家。略：十年征戰（僅六句）'),
    ('木蘭性格', '孝順、堅毅勇敢、不慕名利、柔中帶剛'),
    ('花黃', '古代女子貼於額頭的妝飾，象徵恢復女兒身份'),
    ('結尾比喻', '雄兔撲朔+雌兔迷離→兩兔同跑分不清，比喻木蘭不被識破'),
    ('金柝', '銅製行軍用具，白天煮飯、晚上打更'),
    ('策勳十二轉', '虛數，代表功勳極高，升遷很多次'),
    ('三大主題', '代父從軍（孝）、保家衛國（忠）、不慕名利（義）'),
    ('考試最愛考', '①詳略筆法 ②結尾比喻 ③木蘭人物特質 ④不用尚書郎用意'),
]

lt_hdr = lookup_table.rows[0].cells
lt_hdr[0].text = '關鍵字'
lt_hdr[1].text = '說明'
for cell in lt_hdr:
    set_cell_color(cell, '2E4057')
    for run in cell.paragraphs[0].runs:
        set_font(run, bold=True, color=(255,255,255))

lt_colors = ['F0F8FF','FFFFFF']
for idx, (k, v) in enumerate(lookup_data):
    row = lookup_table.add_row().cells
    row[0].text = k
    row[1].text = v
    bg = lt_colors[idx % 2]
    for cell in row:
        set_cell_color(cell, bg)
    row[0].width = Cm(4)
    row[1].width = Cm(13)

doc.add_paragraph()

# === 複習題 ===
add_heading(doc, '課後複習挑戰（附解答）', level=2, color=(30,130,80))

quiz_data = [
    ('Q1 木蘭決定從軍的原因是什麼？',
     '因為徵兵文書每卷都有父親名字，父親年老且家中無兄長，故女扮男裝代父從軍。\n關鍵句：「阿爺無大兒，木蘭無長兄，願為市鞍馬，從此替爺征。」'),
    ('Q2 本詩「詳略得當」的寫法，哪裡詳寫、哪裡略寫？原因？',
     '詳寫：準備出征、返鄉場景（情感豐富）\n略寫：十年征戰（僅六句帶過）\n原因：主旨在刻畫孝行人物形象，而非戰爭本身，故詳人情略戰事。'),
    ('Q3 「兩兔傍地走，安能辨我是雄雌」使用了什麼修辭？含義？',
     '修辭：比喻\n含義：雄兔、雌兔一起跑時分不清，比喻木蘭女扮男裝與男子無異，突顯巾幗英雄形象，也昇華「性別不礙能力」的主題。'),
    ('Q4 木蘭拒絕「尚書郎」表現了哪些性格特質？',
     '①不慕名利：功成不貪戀官位\n②孝順：心念還鄉陪伴父母\n③真性情：從軍為義，回家為情，首尾一貫'),
    ('Q5 樂府詩與近體詩最大格律差異？',
     '樂府詩：雜言、押韻自由可換韻、無平仄限制\n近體詩：字數固定、韻腳嚴格不可換韻、有平仄規定、律詩必須對仗\n口訣：「樂府自由自在，近體規矩森嚴」'),
]

for q, a in quiz_data:
    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(8)
    qp.paragraph_format.space_after = Pt(2)
    qr = qp.add_run(q)
    set_font(qr, size=11, bold=True, color=(30,80,150))
    set_para_shading(qp, 'E8F4FD')

    ap = doc.add_paragraph()
    ap.paragraph_format.space_after = Pt(6)
    ap.paragraph_format.left_indent = Cm(0.5)
    ar = ap.add_run(a)
    set_font(ar, size=10.5, color=(30,100,50))
    set_para_shading(ap, 'F0FFF4')

# === 儲存 ===
output_path = r'C:/Users/北興/Desktop/claude-ai/木蘭詩_迷因講義.docx'
doc.save(output_path)
print('OK:', output_path)
