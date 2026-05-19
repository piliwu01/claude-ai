import warnings
warnings.filterwarnings("ignore")

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
import threading
import os
import sys
import random
import zipfile
from pathlib import Path

# 確保能 import templates
sys.path.insert(0, os.path.dirname(__file__))
from templates import (
    TEXT_KEYWORDS, generate_puzzle, generate_game
)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ──────────────────────────────────────────
# 色彩
# ──────────────────────────────────────────
C_BG      = "#f0f4f8"
C_CARD    = "#ffffff"
C_BTN_PRI = "#2563eb"
C_BTN_RUN = "#16a34a"
C_BTN_SEC = "#e2e8f0"
C_TITLE   = "#1e293b"
C_SUB     = "#64748b"
C_LOG_BG  = "#1e293b"
C_LOG_FG  = "#94a3b8"
C_OK      = "#4ade80"
C_ERR     = "#f87171"
C_SEP     = "#475569"

FONT_MAIN = "微軟正黑體"
FONT_MONO = "Consolas"

COURSE_LIST = list(TEXT_KEYWORDS.keys())
PUZZLE_TYPES = ["字謎", "數字謎", "密碼謎", "邏輯推理", "圖像描述"]


# ──────────────────────────────────────────
# Word 輸出
# ──────────────────────────────────────────

def _fix_font(doc):
    """修主題 XML 避免 MS Mincho 覆蓋中文字體。"""
    from lxml import etree
    theme_elm = doc.core_properties.element
    # 直接修改 document 的 font theme
    try:
        nsmap = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
        theme_part = doc.part.theme_part
        if theme_part is not None:
            tree = theme_part._element
            for ea in tree.findall('.//a:ea', nsmap):
                ea.set('typeface', '微軟正黑體')
    except Exception:
        pass


def export_to_docx(game_data: dict, output_path: str,
                   show_hints: bool = True, show_answers: bool = False):
    if not DOCX_OK:
        raise RuntimeError("請先安裝 python-docx：pip install python-docx")

    doc = Document()

    # 頁面設定 A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.5)

    def set_font(run, size=12, bold=False, color=None):
        run.font.name = FONT_MAIN
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        set_font(run, size, bold, color)
        return p

    theme = game_data.get("theme", "")
    is_game = "stages" in game_data

    # ── 封面 ──
    add_para("密室逃脫", 24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(30, 41, 59))
    add_para(f"主題：{theme}", 16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(37, 99, 235))
    doc.add_paragraph()

    if is_game:
        # 故事背景
        add_para("【故事背景】", 13, bold=True)
        add_para(game_data["story"], 11)
        doc.add_paragraph()

        for stage in game_data["stages"]:
            sn = stage["stage_num"]
            puzzle = stage["puzzle"]
            link   = stage["link"]

            add_para(f"第 {sn} 關：{puzzle['type']}", 13, bold=True,
                     color=(30, 41, 59))
            add_para(puzzle["question"], 11)
            if show_hints:
                add_para(puzzle["hint"], 10, color=(100, 116, 139))
            if show_answers:
                add_para(f"答案：{puzzle['answer']}", 10, bold=True,
                         color=(22, 163, 74))
            add_para(f"🔗 {link}", 10, color=(71, 85, 105))
            doc.add_paragraph()
    else:
        # 單一謎題
        puzzle = game_data["puzzle"]
        add_para(f"謎題類型：{puzzle['type']}", 13, bold=True)
        add_para(puzzle["question"], 11)
        if show_hints:
            add_para(puzzle["hint"], 10, color=(100, 116, 139))
        if show_answers:
            add_para(f"答案：{puzzle['answer']}", 10, bold=True,
                     color=(22, 163, 74))

    try:
        _fix_font(doc)
    except Exception:
        pass

    doc.save(output_path)


# ──────────────────────────────────────────
# 主視窗
# ──────────────────────────────────────────

class EscapeRoomApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("密室逃脫題目發想器 v1.0")
        self.resizable(True, True)
        self.minsize(600, 700)
        self.configure(bg=C_BG)

        self._game_data = None
        self._build_ui()

    # ── 建立介面 ──────────────────────────

    def _build_ui(self):
        pad = dict(padx=16, pady=6)

        # 大標題
        tk.Label(self, text="密室逃脫題目發想器",
                 font=(FONT_MAIN, 18, "bold"),
                 bg=C_BG, fg=C_TITLE).pack(pady=(18, 2))
        tk.Label(self, text="快速生成課堂密室逃脫謎題",
                 font=(FONT_MAIN, 10),
                 bg=C_BG, fg=C_SUB).pack(pady=(0, 10))

        # ── 模式選擇 ──
        card_mode = self._card("模式選擇")
        self._mode = tk.StringVar(value="single")
        tk.Radiobutton(card_mode, text="單一謎題", variable=self._mode,
                       value="single", bg=C_CARD,
                       font=(FONT_MAIN, 10)).pack(side="left", padx=8)
        tk.Radiobutton(card_mode, text="完整遊戲（3～5 關）", variable=self._mode,
                       value="game", bg=C_CARD,
                       font=(FONT_MAIN, 10)).pack(side="left", padx=8)

        # ── 主題設定 ──
        card_theme = self._card("主題設定")
        self._theme_src = tk.StringVar(value="course")

        row1 = tk.Frame(card_theme, bg=C_CARD)
        row1.pack(fill="x", pady=2)
        tk.Radiobutton(row1, text="課文：", variable=self._theme_src,
                       value="course", bg=C_CARD,
                       font=(FONT_MAIN, 10),
                       command=self._toggle_theme).pack(side="left")
        self._course_var = tk.StringVar(value=COURSE_LIST[0])
        self._course_cb = ttk.Combobox(row1, textvariable=self._course_var,
                                       values=COURSE_LIST, state="readonly",
                                       font=(FONT_MAIN, 10), width=14)
        self._course_cb.pack(side="left", padx=4)

        row2 = tk.Frame(card_theme, bg=C_CARD)
        row2.pack(fill="x", pady=2)
        tk.Radiobutton(row2, text="自訂：", variable=self._theme_src,
                       value="custom", bg=C_CARD,
                       font=(FONT_MAIN, 10),
                       command=self._toggle_theme).pack(side="left")
        self._custom_entry = tk.Entry(row2, font=(FONT_MAIN, 10),
                                      relief="solid", bd=1, width=18,
                                      state="disabled")
        self._custom_entry.pack(side="left", padx=4, ipady=3)

        # ── 謎題類型 ──
        card_type = self._card("謎題類型（可多選）")
        self._type_vars = {}
        for i, pt in enumerate(PUZZLE_TYPES):
            var = tk.BooleanVar(value=(i < 2))
            self._type_vars[pt] = var
            tk.Checkbutton(card_type, text=pt, variable=var,
                           bg=C_CARD, font=(FONT_MAIN, 10)).pack(side="left", padx=6)

        # ── 選項 ──
        card_opt = self._card("輸出選項")
        self._show_hints   = tk.BooleanVar(value=True)
        self._show_answers = tk.BooleanVar(value=False)
        tk.Checkbutton(card_opt, text="顯示提示", variable=self._show_hints,
                       bg=C_CARD, font=(FONT_MAIN, 10)).pack(side="left", padx=8)
        tk.Checkbutton(card_opt, text="顯示答案", variable=self._show_answers,
                       bg=C_CARD, font=(FONT_MAIN, 10)).pack(side="left", padx=8)

        # 若完整遊戲，顯示關卡數
        self._num_stages_label = tk.Label(card_opt, text="關卡數：",
                                          bg=C_CARD, font=(FONT_MAIN, 10))
        self._num_stages_label.pack(side="left", padx=(16, 2))
        self._num_stages = tk.IntVar(value=3)
        tk.Spinbox(card_opt, from_=3, to=5, textvariable=self._num_stages,
                   width=3, font=(FONT_MAIN, 10)).pack(side="left")

        # ── 執行按鈕 ──
        tk.Button(self, text="開始發想",
                  font=(FONT_MAIN, 13, "bold"),
                  bg=C_BTN_RUN, fg="white",
                  relief="flat", padx=30, pady=10, cursor="hand2",
                  command=self._run).pack(pady=10)

        # ── 結果預覽 ──
        frame_log = tk.Frame(self, bg=C_BG)
        frame_log.pack(padx=16, pady=(0, 6), fill="both", expand=True)
        tk.Label(frame_log, text="結果預覽", bg=C_BG, fg=C_SUB,
                 font=(FONT_MAIN, 9)).pack(anchor="w")
        self._log = ScrolledText(frame_log, width=68, height=14,
                                 bg=C_LOG_BG, fg=C_LOG_FG,
                                 font=(FONT_MONO, 9),
                                 relief="flat", state="disabled")
        self._log.pack(fill="both", expand=True)
        self.rowconfigure(0, weight=1)
        frame_log.pack_propagate(False)
        self._log.tag_config("ok",  foreground=C_OK)
        self._log.tag_config("err", foreground=C_ERR)
        self._log.tag_config("sep", foreground=C_SEP)

        # ── 底部按鈕 ──
        btn_row = tk.Frame(self, bg=C_BG)
        btn_row.pack(pady=(0, 16))
        tk.Button(btn_row, text="重新發想",
                  font=(FONT_MAIN, 10), bg=C_BTN_SEC,
                  relief="flat", padx=12, cursor="hand2",
                  command=self._run).pack(side="left", padx=8)
        tk.Button(btn_row, text="匯出 Word",
                  font=(FONT_MAIN, 10, "bold"),
                  bg=C_BTN_PRI, fg="white",
                  relief="flat", padx=12, cursor="hand2",
                  command=self._export).pack(side="left", padx=8)

    def _card(self, title: str) -> tk.Frame:
        lf = tk.LabelFrame(self, text=f"  {title}  ",
                           font=(FONT_MAIN, 10, "bold"),
                           bg=C_CARD, fg="#334155",
                           relief="groove", bd=1,
                           padx=10, pady=8)
        lf.pack(fill="x", padx=16, pady=4)
        return lf

    def _toggle_theme(self):
        if self._theme_src.get() == "course":
            self._course_cb.configure(state="readonly")
            self._custom_entry.configure(state="normal")
            self._custom_entry.configure(state="disabled")
        else:
            self._course_cb.configure(state="disabled")
            self._custom_entry.configure(state="normal")

    # ── 日誌 ──────────────────────────────

    def _log_write(self, msg: str, tag: str = ""):
        self._log.configure(state="normal")
        self._log.insert("end", msg + "\n", tag)
        self._log.see("end")
        self._log.configure(state="disabled")

    def _log_clear(self):
        self._log.configure(state="normal")
        self._log.delete("1.0", "end")
        self._log.configure(state="disabled")

    # ── 取得主題 ──────────────────────────

    def _get_theme(self) -> str:
        if self._theme_src.get() == "course":
            return self._course_var.get()
        else:
            v = self._custom_entry.get().strip()
            return v if v else "未知主題"

    def _get_types(self) -> list:
        return [pt for pt, var in self._type_vars.items() if var.get()]

    # ── 執行 ──────────────────────────────

    def _run(self):
        types = self._get_types()
        if not types:
            messagebox.showwarning("提醒", "請至少選一種謎題類型！")
            return
        threading.Thread(target=self._do_generate, daemon=True).start()

    def _do_generate(self):
        self._log_clear()
        theme = self._get_theme()
        types = self._get_types()
        mode  = self._mode.get()

        self._log_write(f"=== 主題：{theme} | 模式：{'完整遊戲' if mode == 'game' else '單一謎題'} ===", "sep")

        if mode == "game":
            n = self._num_stages.get()
            data = generate_game(theme, types, num_stages=n)
            self._game_data = data

            self._log_write(f"\n📖 故事背景\n{data['story']}\n")
            for stage in data["stages"]:
                sn = stage["stage_num"]
                p  = stage["puzzle"]
                self._log_write(f"── 第 {sn} 關：{p['type']} ──", "sep")
                self._log_write(p["question"])
                if self._show_hints.get():
                    self._log_write(p["hint"])
                if self._show_answers.get():
                    self._log_write(f"答案：{p['answer']}")
                self._log_write(f"🔗 {stage['link']}\n")
        else:
            if not types:
                return
            ptype = random.choice(types)
            puzzle = generate_puzzle(theme, ptype)
            self._game_data = {"theme": theme, "puzzle": puzzle}

            self._log_write(f"── 謎題類型：{puzzle['type']} ──", "sep")
            self._log_write(puzzle["question"])
            if self._show_hints.get():
                self._log_write(puzzle["hint"])
            if self._show_answers.get():
                self._log_write(f"答案：{puzzle['answer']}")

        self._log_write("\n✅ 發想完成！可按「匯出 Word」儲存。", "ok")

    # ── 匯出 ──────────────────────────────

    def _export(self):
        if not self._game_data:
            messagebox.showwarning("提醒", "請先按「開始發想」產生題目！")
            return
        if not DOCX_OK:
            messagebox.showerror("錯誤", "請先安裝 python-docx：\npip install python-docx")
            return

        theme = self._get_theme()
        default_name = f"密室逃脫_{theme}.docx"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word 文件", "*.docx")],
            initialfile=default_name,
            title="儲存 Word 文件",
        )
        if not path:
            return

        try:
            export_to_docx(
                self._game_data, path,
                show_hints=self._show_hints.get(),
                show_answers=self._show_answers.get(),
            )
            self._log_write(f"✅ 已匯出：{path}", "ok")
            messagebox.showinfo("完成", f"Word 文件已儲存：\n{path}")
        except Exception as e:
            self._log_write(f"❌ 匯出失敗：{e}", "err")
            messagebox.showerror("錯誤", str(e))


# ──────────────────────────────────────────
if __name__ == "__main__":
    app = EscapeRoomApp()
    app.mainloop()
